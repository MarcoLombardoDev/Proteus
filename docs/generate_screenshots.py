#!/usr/bin/env python
# Proteus — Rebranding Tool
# Copyright (C) 2026 Marco Lombardo
#
# SPDX-License-Identifier: AGPL-3.0-or-later
# Distributed WITHOUT ANY WARRANTY; see LICENSE for the full terms.
# A commercial licence, without the AGPL's obligations, is available for use
# in proprietary or closed-source products — see COMMERCIAL-LICENSE.md.

"""
Regenerate the README screenshots.

Boots the real application against synthetic sample data and captures one PNG
per tab. Nothing is faked: the windows below are the actual interface.

The script is deliberately side-effect free with respect to the machine it runs
on — the settings file, the log folder and the sample images all live in a
temporary directory that is removed on exit, so running it never touches your
real configuration.

Usage:
    xvfb-run -a python docs/generate_screenshots.py                 # English
    xvfb-run -a python docs/generate_screenshots.py --language it   # Italian
    SHOTDIR=/tmp/shots xvfb-run -a python docs/generate_screenshots.py

Requires `mss` for the screen capture:
    pip install mss
"""

from __future__ import annotations

import argparse
import importlib.util
import os
import shutil
import sys
import tempfile
import time

PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, PROJECT_ROOT)

DEFAULT_SHOTDIR = os.path.join(PROJECT_ROOT, "docs", "screenshots")

#: Colours of the sample logos: the "old" ones being replaced and the "new"
#: ones coming from the source folder. Flat colour blocks keep the previews
#: readable and the repository free of any real brand asset.
OLD_COLOUR = (196, 62, 58)
NEW_COLOUR = (38, 92, 168)

#: One entry per screenshot: (file name, tab index).
SHOTS = (
    ("01_configuration", 0),
    ("02_scan_results", 1),
    ("03_matches", 2),
    ("04_replacement", 3),
    ("05_content_search", 0),
)

#: The command-line capture is not a window: it is the real output of a real
#: run, rendered as a terminal. See `render_terminal`.
CLI_SHOT = "06_command_line"


def build_sample_tree(root: str) -> tuple[str, str]:
    """
    Create a small but representative folder tree.

    It intentionally covers every case the matcher has to grade: exact hits, a
    jpg/jpeg cross-match, an SVG whose size comes from the markup, a weak match
    and a file with no counterpart at all.
    """
    from PIL import Image, ImageDraw

    scan = os.path.join(root, "share")
    source = os.path.join(root, "new_logos")

    def png(folder: str, name: str, size, colour):
        path = os.path.join(folder, name)
        os.makedirs(os.path.dirname(path), exist_ok=True)
        # A mark rather than a flat block: content search needs internal
        # structure to hash, and a flat rectangle has none.
        image = Image.new("RGB", size, (255, 255, 255))
        draw = ImageDraw.Draw(image)
        w, h = size
        draw.ellipse((w * 0.04, h * 0.15, w * 0.32, h * 0.85), fill=colour)
        draw.rectangle((w * 0.38, h * 0.30, w * 0.94, h * 0.48), fill=colour)
        draw.rectangle((w * 0.38, h * 0.56, w * 0.72, h * 0.72), fill=(120, 120, 120))
        image.save(path)

    def svg(folder: str, name: str, width: int, height: int):
        path = os.path.join(folder, name)
        os.makedirs(os.path.dirname(path), exist_ok=True)
        with open(path, "w", encoding="utf-8") as fh:
            fh.write('<svg xmlns="http://www.w3.org/2000/svg" '
                     f'width="{width}px" height="{height}px"/>')

    # Files to be replaced. The last three deliberately carry names no
    # wildcard would guess: they exist to show off content search.
    png(scan, "website/logo_header.png", (240, 80), OLD_COLOUR)
    png(scan, "website/logo_footer.png", (120, 40), OLD_COLOUR)
    png(scan, "intranet/logo_small.png", (64, 64), OLD_COLOUR)   # weak match
    png(scan, "print/logo_press.jpg", (300, 100), OLD_COLOUR)    # jpg -> jpeg
    png(scan, "legacy/logo_orphan.gif", (50, 50), OLD_COLOUR)    # no match
    svg(scan, "web/logo_vector.svg", 500, 200)
    png(scan, "web/header_bg.png", (240, 80), OLD_COLOUR)
    png(scan, "archive/PROGETTO2014.png", (480, 160), OLD_COLOUR)

    # New logos.
    png(source, "logo_header.png", (240, 80), NEW_COLOUR)
    png(source, "logo_footer.png", (120, 40), NEW_COLOUR)
    png(source, "logo_press.jpeg", (300, 100), NEW_COLOUR)
    svg(source, "logo_vector.svg", 500, 200)

    # A Word document with the same logo embedded in it — where a logo usually
    # hides in real life, and the one case a file-name search cannot reach.
    _sample_document(os.path.join(scan, "reports", "annual_report.docx"),
                     os.path.join(scan, "website", "logo_header.png"))

    # A PDF carrying the same logo as a raster image, and one whose artwork is
    # vector-only. The second is there on purpose: it is the case Proteus cannot
    # replace and therefore has to report, and a screenshot of that is worth more
    # than a paragraph claiming it happens.
    _sample_pdf(os.path.join(scan, "print", "brochure.pdf"),
                os.path.join(scan, "website", "logo_header.png"))
    _vector_pdf(os.path.join(scan, "print", "flyer_vector.pdf"))

    # A reference copy of the OLD logo, for the content-search screenshot. It
    # lives outside the scanned tree, as a real one would.
    png(os.path.join(root, "reference"), "old_logo.png", (240, 80), OLD_COLOUR)

    return scan, source


def _sample_pdf(path: str, image: str) -> None:
    """A one-page PDF holding `image` as a raster picture."""
    from PIL import Image

    os.makedirs(os.path.dirname(path), exist_ok=True)
    Image.open(image).convert("RGB").save(path)


def _vector_pdf(path: str) -> bool:
    """
    A PDF whose logo is drawn with path operators, if pypdf is installed.

    Proteus cannot see a vector logo, which is exactly why this file is in the
    sample data: it makes the "needs manual attention" report appear.
    """
    try:
        import pypdf
        from pypdf.generic import DecodedStreamObject, NameObject
    except ImportError:
        print("  (pypdf not installed: skipping the vector PDF sample)")
        return False

    os.makedirs(os.path.dirname(path), exist_ok=True)
    writer = pypdf.PdfWriter()
    page = writer.add_blank_page(240, 80)
    content = DecodedStreamObject()
    content.set_data(b"0.77 0.24 0.23 rg 12 20 90 40 re f")
    page[NameObject("/Contents")] = writer._add_object(content)
    with open(path, "wb") as handle:
        writer.write(handle)
    return True


def _sample_document(path: str, image: str) -> bool:
    """
    Build a small .docx around `image`, if python-docx is installed.

    Optional on purpose: the capture is still useful without it, and the
    library is a test dependency rather than a runtime one. Proteus itself
    reads and rewrites these packages with the standard library alone — the
    library is used here only to produce a document Word would really open.
    """
    try:
        import docx
        from docx.shared import Inches
    except ImportError:
        print("  (python-docx not installed: skipping the Office sample)")
        return False

    os.makedirs(os.path.dirname(path), exist_ok=True)
    document = docx.Document()
    document.add_picture(image, width=Inches(2.4))
    document.add_heading("Annual Report", level=1)
    document.add_paragraph(
        "The logo above is embedded in the document, not linked. A file-name "
        "search cannot see it; Proteus can.")
    document.save(path)
    return True


def capture(path: str, window=None) -> None:
    """
    Grab the virtual screen into `path`, cropped to `window` when given.

    Without the crop the PNG carries whatever unused desktop the virtual screen
    happens to have — on a 1280×1024 Xvfb that is a wide black margin below the
    window, which is both ugly and a waste of the reader's screen.
    """
    import mss
    from PIL import Image

    with mss.mss() as sct:
        shot = sct.grab(sct.monitors[1])
        image = Image.frombytes("RGB", shot.size, shot.bgra, "raw", "BGRX")

    if window is not None:
        box = (window.winfo_rootx(), window.winfo_rooty(),
               window.winfo_rootx() + window.winfo_width(),
               window.winfo_rooty() + window.winfo_height())
        # Clamp: a window may legitimately extend past the virtual screen.
        box = (max(box[0], 0), max(box[1], 0),
               min(box[2], image.width), min(box[3], image.height))
        if box[2] > box[0] and box[3] > box[1]:
            image = image.crop(box)

    image.save(path)


# ---------------------------------------------------------------------------
# The command line
# ---------------------------------------------------------------------------

#: Terminal colours. Dark, because that is what a terminal looks like, and
#: because it distinguishes the CLI shot from the interface ones at a glance.
TERM_BG = (24, 26, 32)
TERM_FG = (208, 212, 220)
TERM_PROMPT = (108, 176, 116)
TERM_COMMAND = (232, 236, 242)
TERM_ALERT = (226, 116, 106)
TERM_DIM = (128, 134, 148)


def _mono_font(size: int):
    """A monospaced face, whatever this machine happens to have."""
    from PIL import ImageFont

    candidates = (
        "/usr/share/fonts/truetype/dejavu/DejaVuSansMono.ttf",
        "/usr/share/fonts/truetype/liberation/LiberationMono-Regular.ttf",
        "/Library/Fonts/Menlo.ttc",
        "C:\\Windows\\Fonts\\consola.ttf",
    )
    for candidate in candidates:
        if os.path.exists(candidate):
            return ImageFont.truetype(candidate, size)
    # Falls back to the bitmap font: ugly, but never crashes the docs build.
    return ImageFont.load_default()


def run_cli(workdir: str, *args: str) -> list[tuple[str, str]]:
    """
    Run the real command line and return its transcript.

    Each entry is (kind, text) where kind is prompt / out / err, so the
    renderer can colour stderr differently — which is the whole point of the
    screenshot: showing that a refusal is loud.
    """
    import subprocess

    shown = " ".join(a.replace(workdir + os.sep, "").replace(workdir, ".")
                     for a in args)
    lines: list[tuple[str, str]] = [("prompt", f"proteus {shown}")]

    result = subprocess.run(
        [sys.executable, os.path.join(PROJECT_ROOT, "main.py"), *args],
        capture_output=True, text=True, cwd=workdir,
    )
    for stream, kind in ((result.stdout, "out"), (result.stderr, "err")):
        for line in stream.splitlines():
            lines.append((kind, line.replace(workdir + os.sep, "")
                                    .replace(workdir, ".")))
    lines.append(("exit", f"$? = {result.returncode}"))
    return lines


def render_terminal(path: str, blocks: list[list[tuple[str, str]]]) -> None:
    """Draw captured transcripts as a terminal window."""
    from PIL import Image, ImageDraw

    font = _mono_font(15)
    bold = _mono_font(15)
    pad, leading = 24, 22

    lines = [line for block in blocks for line in (*block, ("gap", ""))][:-1]
    height = pad * 2 + leading * len(lines)

    probe = Image.new("RGB", (1, 1))
    char = ImageDraw.Draw(probe).textlength("M", font=font) or 8
    # Width from the longest line actually captured. Hardcoding it silently
    # truncated the one line most worth reading — the refusal.
    width = max(len(text) for _kind, text in lines) + 4
    image = Image.new("RGB", (int(pad * 2 + char * width), height), TERM_BG)
    draw = ImageDraw.Draw(image)

    y = pad
    for kind, text in lines:
        if kind == "prompt":
            draw.text((pad, y), "$", font=bold, fill=TERM_PROMPT)
            draw.text((pad + char * 2, y), text, font=bold, fill=TERM_COMMAND)
        elif kind == "err":
            draw.text((pad, y), text, font=font, fill=TERM_ALERT)
        elif kind == "exit":
            draw.text((pad, y), text, font=font, fill=TERM_DIM)
        elif kind != "gap":
            draw.text((pad, y), text, font=font, fill=TERM_FG)
        y += leading

    image.save(path)


def main(argv: list[str]) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--language", default="en",
                        help="interface language to capture (en, it)")
    parser.add_argument("--outdir", default=os.environ.get("SHOTDIR", DEFAULT_SHOTDIR),
                        help="destination folder for the PNG files")
    args = parser.parse_args(argv[1:])

    if not os.environ.get("DISPLAY") and os.name != "nt":
        print("No DISPLAY found. Run this under Xvfb:\n"
              "  xvfb-run -a python docs/generate_screenshots.py", file=sys.stderr)
        return 1

    if importlib.util.find_spec("mss") is None:
        print("The 'mss' package is required: pip install mss", file=sys.stderr)
        return 1

    workdir = tempfile.mkdtemp(prefix="proteus_shots_")
    try:
        import tkinter as tk

        import core
        import i18n

        # Redirect settings and logs into the throwaway folder *before* the
        # application reads them, so the real configuration is left alone.
        core.writable_app_dir = lambda sub: _ensure(os.path.join(workdir, sub))

        scan, source = build_sample_tree(workdir)
        reference = os.path.join(workdir, "reference", "old_logo.png")
        core.save_settings({
            "language": args.language,
            "source_folder": source,
            "scan_folder": scan,
            "search_pattern": "logo*",
            "include_office": True,
            "include_pdf": True,
            "backup": True,
            "dry_run": False,
        })

        from rebranding_tool import RebrandingToolApp

        os.makedirs(args.outdir, exist_ok=True)

        root = tk.Tk()
        app = RebrandingToolApp(root)
        root.update()

        def settle(seconds: float = 0.6) -> None:
            """Pump the event loop, waiting out any running worker."""
            deadline = time.time() + max(seconds, 0.1)
            while app._busy() or time.time() < deadline:
                root.update()
                time.sleep(0.02)
                if time.time() > deadline + 30:
                    break
            for _ in range(6):
                root.update()
                time.sleep(0.03)

        # Walk the wizard exactly as a user would.
        settle()
        _shoot(app, root, args.outdir, "01_configuration", 0, settle)

        app._start_scan()
        settle(1.0)
        _shoot(app, root, args.outdir, "02_scan_results", 1, settle)

        app._start_matching()
        settle(1.0)
        rows = app._match_tree.get_children()
        if rows:
            # Select the weak match so both preview panes are populated.
            app._match_tree.selection_set(rows[0])
        settle()
        _shoot(app, root, args.outdir, "03_matches", 2, settle)

        _shoot(app, root, args.outdir, "04_replacement", 3, settle)

        # Back to tab ①, this time set up for a content search: the mode that
        # finds a logo hiding under a name no wildcard would guess.
        app._search_mode.set("content")
        app._references = [reference]
        app._on_search_mode_changed()
        app._refresh_references_label()
        app.search_pattern.set("")
        settle(0.4)
        _shoot(app, root, args.outdir, "05_content_search", 0, settle)

        print(f"{len(SHOTS)} screenshots written to {args.outdir} "
              f"(language: {i18n.get_language()})")
        root.destroy()

        # The command line has no window: its transcript is captured by
        # running it for real and drawing the output.
        shoot_command_line(workdir, scan, source, reference, args.outdir)
        return 0

    finally:
        shutil.rmtree(workdir, ignore_errors=True)


def shoot_command_line(workdir: str, scan: str, source: str,
                       reference: str, outdir: str) -> None:
    """
    Capture three real runs: a dry run, a refusal, and the applied campaign.

    The middle one is the reason this screenshot exists. A refusal has to be
    legible in a log nobody reads until something has gone wrong, so it is
    worth showing what one actually looks like.
    """
    common = ["--scan", scan, "--source", source]
    blocks = [
        run_cli(workdir, *common, "--pattern", "logo*.png", "--verbose"),
        run_cli(workdir, *common, "--reference", reference,
                "--similarity", "70", "--office", "--apply"),
        # Names the PDFs explicitly, which is what makes a vector-only file a
        # finding rather than noise — and shows exit code 5.
        run_cli(workdir, *common, "--pattern", "*.pdf", "--pdf", "--apply"),
    ]
    path = os.path.join(outdir, f"{CLI_SHOT}.png")
    render_terminal(path, blocks)
    print(f"  captured {CLI_SHOT}.png")


def _ensure(path: str) -> str:
    os.makedirs(path, exist_ok=True)
    return path


def _shoot(app, root, outdir: str, name: str, tab: int, settle) -> None:
    app.notebook.select(tab)
    settle(0.4)
    capture(os.path.join(outdir, f"{name}.png"), window=root)
    print(f"  captured {name}.png")


if __name__ == "__main__":
    sys.exit(main(sys.argv))
